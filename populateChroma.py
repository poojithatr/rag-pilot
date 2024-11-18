import argparse
import os
import shutil
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema.document import Document
from starter import get_embedding_function
from langchain_community.vectorstores import Chroma
from langchain_ollama import OllamaEmbeddings
from docx import Document as DocxDocument
from langchain.schema import Document
import umap
import plotly.express as px
from sklearn.preprocessing import StandardScaler
import numpy as np

CHROMA_PATH = "chroma"
DATA_PATH = "data"

def main():
    # Check if the database should be cleared (using the --reset flag).
    parser = argparse.ArgumentParser()
    parser.add_argument("--reset", action="store_true", help="Reset the database.")
    args = parser.parse_args()
    print("✨ Clearing Database")
    clear_database()

    # Create (or update) the data store.
    documents = load_documents()
    chunks = split_documents(documents)

    # Add documents to Chroma and get the metadata
    document_metadata = add_to_chroma(chunks)

    # Retrieve embeddings from Chroma (this function should fetch embeddings from the database)
    embeddings = get_embeddings_from_chroma()

    # Visualize embeddings with the document metadata
    visualize_embeddings(embeddings, document_metadata, n_components=2, interactive=True)


def get_embeddings_from_chroma():
    db = Chroma(persist_directory=CHROMA_PATH, embedding_function=get_embedding_function())

    # Use the 'query' method to retrieve the embeddings. You may need to adjust based on how Chroma stores embeddings.
    # For example, if you are fetching all embeddings, you could use the following:

    embeddings = db.get(include=["embeddings"])["embeddings"]  # Make sure to specify "embeddings" to retrieve them.

    return embeddings


def load_documents():
    documents = []
    for filename in os.listdir(DATA_PATH):
        doc_path = os.path.join(DATA_PATH, filename)

        if filename.endswith(".txt"):
            with open(doc_path, "r") as file:
                text = file.read()

            # Split the text into paragraphs based on newline characters
            paragraphs = text.split("\n\n")  # assuming paragraphs are separated by blank lines

            # Wrap the text in a Document object
            for para in paragraphs:
                documents.append(Document(page_content=para, metadata={"source": filename}))

        elif filename.endswith(".docx"):
            # Use python-docx to read .docx files
            doc = DocxDocument(doc_path)
            for para in doc.paragraphs:
                # Add each paragraph to the documents list
                documents.append(Document(page_content=para.text, metadata={"source": filename}))

    return documents

def split_documents(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=80,
        length_function=len,
        is_separator_regex=False,
    )
    return text_splitter.split_documents(documents)


def add_to_chroma(chunks: list[Document]):
    # Load the existing database.
    db = Chroma(
        persist_directory=CHROMA_PATH, embedding_function=get_embedding_function()
    )

    # Calculate Page IDs.
    chunks_with_ids = calculate_chunk_ids(chunks)

    # Add or Update the documents.
    existing_items = db.get(include=[])  # IDs are always included by default
    existing_ids = set(existing_items["ids"])
    print(f"Number of existing documents in DB: {len(existing_ids)}")

    # Only add documents that don't exist in the DB.
    new_chunks = []
    document_metadata = []  # Store document metadata here
    for chunk in chunks_with_ids:
        if chunk.metadata["id"] not in existing_ids:
            new_chunks.append(chunk)
            document_metadata.append(chunk.metadata)  # Store metadata for visualization

    if len(new_chunks):
        print(f"👉 Adding new documents: {len(new_chunks)}")
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)
        db.persist()

    return document_metadata  # Return the document metadata

def calculate_chunk_ids(chunks):

    # This will create IDs like "data/monopoly.pdf:6:2"
    # Page Source : Page Number : Chunk Index

    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        current_page_id = f"{source}:{page}"

        # If the page ID is the same as the last one, increment the index.
        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        # Calculate the chunk ID.
        chunk_id = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id

        # Add it to the page meta-data.
        chunk.metadata["id"] = chunk_id

    return chunks


def clear_database():
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)

def visualize_embeddings(embeddings, document_metadata, n_components=2, interactive=True):
    """
    Visualizes embeddings using UMAP and Plotly.
    """
    # Step 1: Ensure embeddings and document_metadata have the same length
    if len(embeddings) != len(document_metadata):
        print(f"Warning: Number of embeddings ({len(embeddings)}) does not match number of document_metadata ({len(document_metadata)})")
        # Handle the mismatch as per your logic, e.g., truncate, pad, or raise an error
        # For example, truncating the longer list to match the shorter one:
        min_len = min(len(embeddings), len(document_metadata))
        embeddings = embeddings[:min_len]
        document_metadata = document_metadata[:min_len]

    # Step 2: Normalize the embeddings
    scaler = StandardScaler()
    embeddings_scaled = scaler.fit_transform(embeddings)

    # Step 3: Reduce the dimensionality of the embeddings using UMAP
    umap_model = umap.UMAP(n_components=n_components)
    umap_embeddings = umap_model.fit_transform(embeddings_scaled)

    # Step 4: Prepare document labels for visualization
    document_labels = [meta.get("source", "Unknown") for meta in document_metadata]

    # Step 5: Visualize using Plotly or Matplotlib
    if interactive:
        if n_components == 2:
            fig = px.scatter(
                x=umap_embeddings[:, 0],
                y=umap_embeddings[:, 1],
                title="UMAP 2D Visualization",
                hover_data={"text": document_labels}
            )
        elif n_components == 3:
            fig = px.scatter_3d(
                x=umap_embeddings[:, 0],
                y=umap_embeddings[:, 1],
                z=umap_embeddings[:, 2],
                title="UMAP 3D Visualization",
                hover_data={"text": document_labels}
            )
        fig.show()


if __name__ == "__main__":
    main()